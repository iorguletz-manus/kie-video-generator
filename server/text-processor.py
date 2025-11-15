#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
COMPLETE AD TEXT PROCESSOR
Combines both phases:
1. Process text with diacritics, overlap, and red on Line 2
2. Add red on Line 1 for overlap pairs

Usage: python3 complete_processor.py
Input: /home/ubuntu/upload/1MANUS-fullad.docx
Output: /home/ubuntu/COMPLETE_OUTPUT.docx
"""

from docx import Document
from docx.shared import RGBColor, Pt
import re
import random

# ============================================================================
# PHASE 1: TEXT PROCESSING FUNCTIONS
# ============================================================================

def add_diacritics(text):
    """Add Romanian diacritics"""
    replacements = {
        'saturat': 'săturat', 'traiasca': 'trăiască', 'traiesc': 'trăiesc', 'traiesti': 'trăiești', 'traiam': 'trăiam',
        'sa': 'să', 'si': 'și', 'asa': 'așa', 'iti': 'îți', 'isi': 'își', 'imi': 'îmi',
        'luna': 'lună', 'viata': 'viața', 'vietii': 'vieții',
        'straduiesc': 'străduiesc', 'reusesc': 'reuşesc', 'reusea': 'reuşea', 'reușeam': 'reuşeam',
        'stiu': 'știu', 'stia': 'știa', 'simti': 'simți', 'poti': 'poți',
        'vina': 'vină', 'renunti': 'renunți', 'pana': 'până',
        'daca': 'dacă', 'supravietuire': 'supraviețuire', 'asculta-ma': 'ascultă-mă',
        'muncesti': 'muncești', 'platesti': 'plătești', 'bucuri': 'bucuri',
        'fura': 'fură', 'linistea': 'liniștea',
        'uiti': 'uiți', 'uita': 'uită', 'platesc': 'plătesc',
        'iesire': 'ieșire',
        'lupta': 'luptă',
        'fara': 'fără', 'plateam': 'plăteam', 'rata': 'rată', 'apareau': 'apăreau',
        'parca': 'parcă', 'acopar': 'acopăr', 'gaurile': 'găurile', 'seara': 'seară',
        'intrebam': 'întrebam', 'chinui': 'chinui', 'ma': 'mă',
        'uitase': 'uitase', 'libertatea': 'libertatea', 'supravietuiasca': 'supraviețuiască',
        'intr-o': 'într-o', 'blocata': 'blocată', 'castig': 'câștig',
        'prinsa': 'prinsă', 'ganduri': 'gânduri',
        'aratat': 'arătat', 'gandurile': 'gândurile', 'fricile': 'fricile',
        'blocheaza': 'blochează', 'sterse': 'șterse', 'metoda': 'metodă', 'mintii': 'minții',
        'sceptica': 'sceptică', 'inceput': 'început', 'deschid': 'deschid',
        'ajuns': 'ajuns', 'visam': 'visam', 'linistit': 'liniștit',
        'astazi': 'astăzi', 'supravietuieste': 'supraviețuiește', 'siguranta': 'siguranță',
        'doresti': 'dorești', 'castigi': 'câștigi', 'masura': 'măsura',
        'recomand': 'recomand', 'toata': 'toată', 'inima': 'inimă', 'aceasta': 'această',
        'prezentata': 'prezentată', 'catre': 'către', 'rescrie': 'rescrie', 'povestea': 'povestea',
        'capat': 'capăt', 'cauza': 'cauză', 'adevarat': 'adevărat', 'cand': 'când', 'ca': 'că',
        'doua': 'două', 'in': 'în',
    }
    
    words = text.split()
    result = []
    for word in words:
        match = re.match(r'^([^\w]*)(.+?)([^\w]*)$', word, re.UNICODE)
        if match:
            prefix, core, suffix = match.groups()
            core_lower = core.lower()
            if core_lower in replacements:
                replacement = replacements[core_lower]
                if core and core[0].isupper():
                    replacement = replacement[0].upper() + replacement[1:]
                result.append(prefix + replacement + suffix)
            else:
                result.append(word)
        else:
            result.append(word)
    
    return ' '.join(result)

def split_into_sentences(text):
    """Split text into sentences based on . ! ?"""
    sentences = []
    current = []
    words = text.split()
    
    for word in words:
        current.append(word)
        # Only split on . ! ? not on :
        if word[-1] in '.!?':
            sentences.append(' '.join(current).strip())
            current = []
    
    if current:
        sentences.append(' '.join(current).strip())
    
    return sentences

def process_short_text(text, min_c=118, max_c=125):
    """
    Handle text < 118 chars.
    Add WORDS from beginning to reach random target between 118-125.
    """
    target = random.randint(min_c, max_c)
    needed = target - len(text)
    
    if needed <= 0:
        return (text, -1, -1, len(text))
    
    # Add WORDS from beginning
    words = text.split()
    added_words = []
    
    # Keep adding words until we reach target
    for _ in range(10):  # Max 10 cycles
        for word in words:
            test_len = len(text + " " + " ".join(added_words + [word]))
            if test_len <= max_c:
                added_words.append(word)
                if test_len >= min_c:
                    # Check if we should continue to get closer to target
                    if test_len >= target:
                        break
                else:
                    continue
            else:
                break
        
        if len(text + " " + " ".join(added_words)) >= min_c:
            break
    
    added_text = " ".join(added_words)
    full_text = text + " " + added_text
    
    # Trim if over max_c (word by word)
    while len(full_text) > max_c and " " in added_text:
        added_words = added_words[:-1]
        added_text = " ".join(added_words)
        full_text = text + " " + added_text
    
    red_start = len(text) + 1
    red_end = len(full_text)
    
    return (full_text, red_start, red_end, len(full_text))

def process_long_sentence_with_overlap(text, min_c=118, max_c=125):
    """
    Handle single sentence > 125 chars with strategic overlap.
    
    Creates 2 versions:
    - Version 1: First N chars (random 118-125), no red
    - Version 2: Last M chars (random 118-125), red on strategic CUT point
    """
    # Version 1: First part
    rand1 = random.randint(min_c, max_c)
    version1 = text[:rand1].rstrip()
    
    # Adjust to end at word boundary
    if ' ' in version1:
        words = version1.split()
        version1 = ' '.join(words)
    
    # Version 2: Last part
    rand2 = random.randint(min_c, max_c)
    version2 = text[-rand2:].lstrip()
    
    # Adjust to start at word boundary
    if ' ' in version2:
        words = version2.split()
        version2 = ' '.join(words)
    
    results = []
    
    # Add version 1 (no red)
    results.append((version1, -1, -1, len(version1)))
    
    # Find strategic CUT point in version2
    # Look for: punctuation (:, ,), transition words, or idea change
    words = version2.split()
    strategic_cut_idx = -1
    
    # Look for punctuation marks that indicate idea change
    for i, word in enumerate(words):
        # Check if word ends with : or ,
        if word[-1] in ':,':
            # Check if remaining text is long enough (>50 chars, ~3+ seconds)
            remaining = ' '.join(words[i+1:])
            if len(remaining) >= 50:
                strategic_cut_idx = i + 1
                break
    
    # If no good punctuation found, look for transition words
    if strategic_cut_idx == -1:
        transition_words = ['dar', 'și', 'iar', 'pentru', 'astfel', 'când', 'dacă']
        for i, word in enumerate(words):
            word_clean = word.lower().strip('.,!?:;')
            if word_clean in transition_words:
                remaining = ' '.join(words[i:])
                if len(remaining) >= 50:
                    strategic_cut_idx = i
                    break
    
    # If still not found, use 30% of text as fallback
    if strategic_cut_idx == -1:
        target_len = int(len(version2) * 0.3)
        current_len = 0
        for i, word in enumerate(words):
            current_len += len(word) + 1
            if current_len >= target_len:
                remaining = ' '.join(words[i+1:])
                if len(remaining) >= 50:
                    strategic_cut_idx = i + 1
                    break
    
    # Apply red marking
    if strategic_cut_idx > 0:
        red_text = ' '.join(words[:strategic_cut_idx])
        red_end = len(red_text)
        results.append((version2, 0, red_end, len(version2)))
    else:
        # Fallback: no red
        results.append((version2, -1, -1, len(version2)))
    
    return results

def process_text(text, min_c=118, max_c=125):
    """
    Main processing function following exact rules.
    """
    text = text.strip()
    if not text:
        return []
    
    # Apply diacritics (DISABLED - uncomment to enable)
    # text = add_diacritics(text)
    length = len(text)
    
    # Rule 1: If 118-125, keep as is
    if min_c <= length <= max_c:
        return [(text, -1, -1, length)]
    
    # Rule 2: If < 118, add from beginning
    elif length < min_c:
        return [process_short_text(text, min_c, max_c)]
    
    # Rule 3: If > 125
    else:
        sentences = split_into_sentences(text)
        
        # Single sentence > 125 - use overlap strategy
        if len(sentences) == 1:
            return process_long_sentence_with_overlap(text, min_c, max_c)
        
        # Multiple sentences - process sequentially
        results = []
        i = 0
        
        while i < len(sentences):
            # Try combining 3 sentences
            if i + 2 < len(sentences):
                combined3 = " ".join(sentences[i:i+3])
                if min_c <= len(combined3) <= max_c:
                    results.append((combined3, -1, -1, len(combined3)))
                    i += 3
                    continue
                elif len(combined3) < min_c:
                    result = process_short_text(combined3, min_c, max_c)
                    results.append(result)
                    i += 3
                    continue
            
            # Try combining 2 sentences
            if i + 1 < len(sentences):
                combined2 = " ".join(sentences[i:i+2])
                if min_c <= len(combined2) <= max_c:
                    results.append((combined2, -1, -1, len(combined2)))
                    i += 2
                    continue
                elif len(combined2) < min_c:
                    result = process_short_text(combined2, min_c, max_c)
                    results.append(result)
                    i += 2
                    continue
            
            # Process single sentence
            results.extend(process_text(sentences[i], min_c, max_c))
            i += 1
        
        return results

# ============================================================================
# PHASE 1: MAIN PROCESSING
# ============================================================================

print("=" * 60)
print("PHASE 1: Processing text with diacritics and overlap")
print("=" * 60)

# Read document
input_doc = Document('/home/ubuntu/upload/1MANUS-fullad.docx')

all_text = []
for para in input_doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

text_content = '\n'.join(all_text)

labels = ['HOOKS:', 'H1:', 'H2:', 'H3:', 'H4:', 'H5:', 'H6:', 'H7:', 'H8:', 'H9:', 
          'MIRROR1', 'DCS & IDENTITY1', 'TRANZITIE1', 'NEW CAUSE1', 'MECHANISM1', 
          'EMOTIONAL PROOF1', 'TRANSFORMATION1', 'CTA1']

lines = text_content.split('\n')
output_data = []
current_text = []

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    is_label = any(line == lbl or line.startswith(lbl) for lbl in labels)
    
    if is_label:
        if current_text:
            full = ' '.join(current_text)
            processed = process_text(full)
            
            for full_text, red_start, red_end, char_count in processed:
                output_data.append({
                    'type': 'text',
                    'text': full_text,
                    'red_start': red_start,
                    'red_end': red_end,
                    'char_count': char_count
                })
            
            current_text = []
        
        output_data.append({'type': 'label', 'text': line})
    else:
        current_text.append(line)

if current_text:
    full = ' '.join(current_text)
    processed = process_text(full)
    
    for full_text, red_start, red_end, char_count in processed:
        output_data.append({
            'type': 'text',
            'text': full_text,
            'red_start': red_start,
            'red_end': red_end,
            'char_count': char_count
        })

# Create document
output_doc = Document()

for item in output_data:
    if item['type'] == 'label':
        p = output_doc.add_paragraph()
        run = p.add_run(item['text'])
        run.bold = True
        run.font.size = Pt(11)
    else:
        p = output_doc.add_paragraph()
        text = item['text']
        red_start = item['red_start']
        red_end = item['red_end']
        char_count = item['char_count']
        
        if red_start >= 0 and red_end >= 0:
            if red_start > 0:
                run = p.add_run(text[:red_start])
            run = p.add_run(text[red_start:red_end])
            run.font.color.rgb = RGBColor(255, 0, 0)
            if red_end < len(text):
                run = p.add_run(text[red_end:])
        else:
            run = p.add_run(text)
        
        run = p.add_run(f" - ")
        run = p.add_run(f"{char_count} chars")
        run.bold = True

# Save intermediate document
output_doc.save('/home/ubuntu/TEMP_PHASE1.docx')

# Verify
issues = []
for i, item in enumerate(output_data):
    if item['type'] == 'text':
        if item['char_count'] < 118 or item['char_count'] > 125:
            issues.append((i, item['char_count']))

print(f"✓ Generated {len([x for x in output_data if x['type'] == 'text'])} text lines")

if issues:
    print(f"⚠ {len(issues)} lines outside 118-125:")
    for idx, count in issues[:5]:
        print(f"  Line {idx}: {count} chars")
else:
    print("✓✓✓ PERFECT! ALL LINES 118-125 CHARS!")

# ============================================================================
# PHASE 2: ADD RED ON LINE 1
# ============================================================================

print("\n" + "=" * 60)
print("PHASE 2: Adding red on Line 1 for overlap pairs")
print("=" * 60)

# Load the document from Phase 1
doc = Document('/home/ubuntu/TEMP_PHASE1.docx')
paragraphs = list(doc.paragraphs)

processed = 0
i = 0

while i < len(paragraphs):
    para = paragraphs[i]
    text = para.text.strip()
    
    # Look for consecutive lines with "chars"
    if 'chars' in text:
        if i + 1 < len(paragraphs) and 'chars' in paragraphs[i + 1].text:
            line1 = para
            line2 = paragraphs[i + 1]
            
            # Check if both have red
            l1_has_red = any(r.font.color and r.font.color.rgb and r.font.color.rgb == RGBColor(255, 0, 0) for r in line1.runs)
            l2_has_red = any(r.font.color and r.font.color.rgb and r.font.color.rgb == RGBColor(255, 0, 0) for r in line2.runs)
            
            if l2_has_red and not l1_has_red:  # L2 has red, L1 doesn't
                # Find first normal run after red on Line 2
                found_red = False
                l2_normal_text = None
                
                for r in line2.runs:
                    is_red = r.font.color and r.font.color.rgb and r.font.color.rgb == RGBColor(255, 0, 0)
                    is_bold = r.bold
                    
                    if is_red:
                        found_red = True
                    elif found_red and not is_bold and len(r.text) > 5:
                        l2_normal_text = r.text
                        break
                
                if l2_normal_text and len(l2_normal_text) >= 5:
                    # Take first 5 chars
                    search_str = l2_normal_text[:5]
                    
                    # Get L1 text
                    l1_text = line1.text
                    if ' - ' in l1_text:
                        l1_text = l1_text.split(' - ')[0].strip()
                    
                    # Search in L1
                    idx = l1_text.find(search_str)
                    
                    if idx != -1:
                        # IMPORTANT: Ensure normal part (before red) is minimum 40 chars
                        # Search backward for natural break point
                        if idx < 40:
                            # Too short! Search backward for punctuation or conjunctions
                            natural_breaks = ['. ', '! ', '? ', ': ', ', ', ' și ', ' dar ', ' pentru că ', ' când ', ' dacă ']
                            best_idx = idx
                            
                            for break_str in natural_breaks:
                                # Find last occurrence before current idx that gives >= 40 chars
                                temp_idx = l1_text.rfind(break_str, 0, idx)
                                if temp_idx != -1:
                                    # Check if this gives >= 40 chars normal part
                                    potential_normal_len = temp_idx + len(break_str)
                                    if potential_normal_len >= 40:
                                        best_idx = potential_normal_len
                                        break
                            
                            idx = best_idx
                    
                    if idx >= 40:
                        # Get char count
                        char_count = ""
                        for r in line1.runs:
                            if r.bold:
                                char_count = r.text
                                break
                        
                        # Rebuild Line 1
                        line1.clear()
                        
                        # Before match (normal)
                        if idx > 0:
                            line1.add_run(l1_text[:idx])
                        
                        # From match to end (red)
                        red_run = line1.add_run(l1_text[idx:])
                        red_run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        # Add char count
                        line1.add_run(" - ")
                        bold_run = line1.add_run(char_count)
                        bold_run.bold = True
                        
                        processed += 1
            
            i += 2
            continue
    
    i += 1

# Save final document
doc.save('/home/ubuntu/COMPLETE_OUTPUT.docx')

print(f"✓ Processed {processed} overlap pairs with red on Line 1")

print("\n" + "=" * 60)
print("COMPLETE! Final output: /home/ubuntu/COMPLETE_OUTPUT.docx")
print("=" * 60)
